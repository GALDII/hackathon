import os
import json
from groq import Groq
from dotenv import load_dotenv
import pypdf
import pandas as pd
from docx import Document
import tiktoken
from typing import List, Dict
import numpy as np

# Load environment variables from .env file
load_dotenv()

# Initialize the Groq client
try:
    client = Groq(
        api_key=os.environ.get("GROQ_API_KEY")
    )
except Exception as e:
    print(f"Error initializing Groq client: {e}")
    client = None

def count_tokens(text: str, model: str = "gpt-3.5-turbo") -> int:
    """Count tokens in text using tiktoken (approximate for Groq models)"""
    try:
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except:
        # Fallback: rough estimation (4 chars per token)
        return len(text) // 4

def chunk_text(text: str, max_tokens: int = 6000) -> List[str]:
    """Split text into chunks that fit within token limits"""
    # Split by pages/sections first (look for common PDF patterns)
    paragraphs = text.split('\n\n')
    
    chunks = []
    current_chunk = ""
    
    for paragraph in paragraphs:
        # Check if adding this paragraph would exceed the limit
        potential_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
        
        if count_tokens(potential_chunk) <= max_tokens:
            current_chunk = potential_chunk
        else:
            # Save current chunk if it has content
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            # Start new chunk with current paragraph
            if count_tokens(paragraph) <= max_tokens:
                current_chunk = paragraph
            else:
                # If single paragraph is too long, split it further
                words = paragraph.split()
                temp_chunk = ""
                for word in words:
                    if count_tokens(temp_chunk + " " + word) <= max_tokens:
                        temp_chunk += " " + word if temp_chunk else word
                    else:
                        if temp_chunk:
                            chunks.append(temp_chunk.strip())
                        temp_chunk = word
                
                current_chunk = temp_chunk
    
    # Add final chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def extract_text_from_pdf(pdf_path):
    """Extracts text from a given PDF file."""
    print(f"Reading PDF from: {pdf_path}")
    with open(pdf_path, 'rb') as file:
        pdf_reader = pypdf.PdfReader(file)
        text = ""
        total_pages = len(pdf_reader.pages)
        print(f"PDF has {total_pages} pages")
        
        for i, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            text += f"\n--- Page {i} ---\n" + page_text
            if i % 10 == 0:  # Progress indicator for large PDFs
                print(f"Processed {i}/{total_pages} pages...")
        
    print("Successfully extracted text from PDF.")
    print(f"Total text length: {len(text)} characters")
    return text

def analyze_excel_file(excel_path):
    """Analyze Excel file and extract insights"""
    print(f"Reading Excel file from: {excel_path}")
    
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(excel_path)
        all_sheets_data = {}
        
        print(f"Excel file has {len(excel_file.sheet_names)} sheet(s): {excel_file.sheet_names}")
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            all_sheets_data[sheet_name] = df
            print(f"Sheet '{sheet_name}': {df.shape[0]} rows, {df.shape[1]} columns")
        
        # Generate comprehensive data summary
        data_summary = generate_excel_summary(all_sheets_data)
        
        return data_summary
        
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None

def generate_excel_summary(all_sheets_data):
    """Generate a comprehensive summary of Excel data"""
    summary_text = "EXCEL DATA ANALYSIS SUMMARY\n" + "="*50 + "\n\n"
    
    for sheet_name, df in all_sheets_data.items():
        summary_text += f"SHEET: {sheet_name}\n" + "-"*30 + "\n"
        summary_text += f"Dimensions: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
        
        # Column information
        summary_text += "COLUMNS:\n"
        for col in df.columns:
            dtype = str(df[col].dtype)
            non_null = df[col].count()
            null_count = df[col].isnull().sum()
            summary_text += f"- {col} ({dtype}): {non_null} non-null, {null_count} null\n"
        
        # Numerical analysis
        numerical_cols = df.select_dtypes(include=[np.number]).columns
        if len(numerical_cols) > 0:
            summary_text += f"\nNUMERICAL ANALYSIS:\n"
            for col in numerical_cols:
                if df[col].count() > 0:  # Only if there's data
                    stats = df[col].describe()
                    summary_text += f"- {col}: Min={stats['min']:.2f}, Max={stats['max']:.2f}, Mean={stats['mean']:.2f}, Std={stats['std']:.2f}\n"
        
        # Categorical analysis
        categorical_cols = df.select_dtypes(include=['object', 'category']).columns
        if len(categorical_cols) > 0:
            summary_text += f"\nCATEGORICAL ANALYSIS:\n"
            for col in categorical_cols[:5]:  # Limit to first 5 categorical columns
                if df[col].count() > 0:
                    unique_count = df[col].nunique()
                    top_values = df[col].value_counts().head(3)
                    summary_text += f"- {col}: {unique_count} unique values. Top: {dict(top_values)}\n"
        
        # Sample data
        summary_text += f"\nSAMPLE DATA (first 3 rows):\n"
        sample_data = df.head(3).to_string(index=False, max_cols=10)
        summary_text += sample_data + "\n\n"
        
        # Key observations
        summary_text += f"KEY OBSERVATIONS:\n"
        summary_text += f"- Total records: {len(df):,}\n"
        summary_text += f"- Complete rows (no nulls): {df.dropna().shape[0]:,}\n"
        summary_text += f"- Numerical columns: {len(numerical_cols)}\n"
        summary_text += f"- Text/categorical columns: {len(categorical_cols)}\n"
        
        if len(numerical_cols) > 0:
            # Find potential KPIs
            high_value_cols = []
            for col in numerical_cols:
                if df[col].count() > 0:
                    max_val = df[col].max()
                    if max_val > 1000:  # Potential financial/business metrics
                        high_value_cols.append((col, max_val))
            
            if high_value_cols:
                summary_text += f"- Potential business metrics found in: {[col for col, _ in high_value_cols]}\n"
        
        summary_text += "\n" + "="*60 + "\n\n"
    
    return summary_text

def analyze_chunk(chunk: str, chunk_index: int) -> Dict:
    """Analyze a single chunk of text"""
    system_prompt = f"""
    You are a highly skilled business analyst AI. You are analyzing chunk {chunk_index + 1} of a larger document.
    
    Extract key business metrics, financial data, and performance indicators from this text chunk.
    Focus on quantitative data, percentages, monetary values, dates, and business KPIs.
    
    Return ONLY a JSON object with this structure:
    {{
      "chunk_summary": "Brief summary of what this chunk contains",
      "metrics_found": [
        {{
          "metric": "Metric name",
          "value": "Value with units",
          "context": "Brief context about where this metric was found"
        }}
      ]
    }}
    
    If no meaningful business metrics are found in this chunk, return an empty metrics_found array.
    """
    
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": chunk}
            ],
            temperature=0.1
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"Error analyzing chunk {chunk_index + 1}: {e}")
        return {"chunk_summary": f"Error processing chunk {chunk_index + 1}", "metrics_found": []}

def synthesize_final_report(chunk_results: List[Dict], total_chunks: int) -> Dict:
    """Synthesize all chunk results into a final comprehensive report"""
    
    # Prepare summary of all findings
    all_metrics = []
    chunk_summaries = []
    
    for i, result in enumerate(chunk_results):
        if result.get("chunk_summary"):
            chunk_summaries.append(f"Chunk {i+1}: {result['chunk_summary']}")
        
        for metric in result.get("metrics_found", []):
            all_metrics.append(metric)
    
    # Create synthesis prompt
    synthesis_prompt = f"""
    You are a senior business analyst creating a comprehensive report from analysis of {total_chunks} document chunks.
    
    Based on the extracted metrics and summaries below, create a final executive report.
    
    Chunk Summaries:
    {chr(10).join(chunk_summaries)}
    
    All Extracted Metrics:
    {json.dumps(all_metrics, indent=2)}
    
    Create a final report with this EXACT JSON structure:
    {{
      "executive_summary": "A comprehensive 3-4 sentence executive summary of the entire document's key findings and business performance",
      "key_metrics": [
        {{
          "metric": "Most important metric name",
          "value": "Value with proper formatting",
          "commentary": "Business significance and context of this metric"
        }}
      ]
    }}
    
    Select the 8-10 MOST IMPORTANT and UNIQUE metrics. Avoid duplicates.
    Prioritize financial metrics, growth rates, and key performance indicators.
    """
    
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": synthesis_prompt}
            ],
            temperature=0.1
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"Error in final synthesis: {e}")
        # Fallback: create report from raw metrics
        return {
            "executive_summary": "Document analysis completed with chunked processing due to document size.",
            "key_metrics": all_metrics[:10]  # Take first 10 metrics as fallback
        }

def get_excel_insights_from_llm(excel_summary):
    """Get insights specifically from Excel data"""
    if not client:
        raise ConnectionError("Groq client not initialized. Check your API key.")

    system_prompt = """
    You are a highly skilled data analyst AI specializing in Excel/spreadsheet analysis. Your task is to analyze the provided Excel data summary and extract key insights, patterns, and business metrics.

    Please provide the output in a structured JSON format ONLY. Do not include any introductory text, explanations, or markdown formatting outside of the JSON structure itself.

    The JSON structure must be as follows:
    {
      "executive_summary": "A concise, 2-3 sentence summary of the overall data insights and key findings from the Excel analysis.",
      "key_metrics": [
        {
          "metric": "Data insight or key finding (e.g., 'Average Revenue per Customer', 'Data Completeness Rate', 'Top Performance Category')",
          "value": "The value or finding (e.g., '$2,450', '95%', 'Product Category A')",
          "commentary": "A brief, one-sentence comment on the business significance of this finding."
        }
      ]
    }

    Focus on:
    - Data quality and completeness
    - Statistical insights from numerical data
    - Key patterns in categorical data
    - Business-relevant metrics that can be derived
    - Data structure and organization insights
    
    Extract at least 6-8 key insights from the Excel data.
    """

    print("Sending Excel summary to Groq LLM for analysis...")
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Here is the Excel data summary:\n\n{excel_summary}"}
            ]
        )
        json_response = response.choices[0].message.content
        print("Received structured JSON data from Groq LLM for Excel analysis.")
        return json.loads(json_response)
    except Exception as e:
        print(f"An error occurred while communicating with the Groq API: {e}")
        return None

def get_insights_from_llm(text_content):
    """
    Handles both small and large documents by chunking when necessary.
    """
    if not client:
        raise ConnectionError("Groq client not initialized. Check your API key.")

    # Check if document is small enough to process directly
    token_count = count_tokens(text_content)
    print(f"Document token count: approximately {token_count:,} tokens")
    
    # If document is small enough (under 8000 tokens), process normally
    if token_count <= 8000:
        print("Document is small enough for direct processing")
        return get_insights_direct(text_content)
    
    # For large documents, use chunking strategy
    print("Document is large - using chunking strategy")
    return get_insights_chunked(text_content)

def get_insights_direct(text_content):
    """Original processing for smaller documents"""
    system_prompt = """
    You are a highly skilled financial analyst AI. Your task is to analyze the provided text from a business document and extract key performance indicators (KPIs), metrics, and a summary.

    Please provide the output in a structured JSON format ONLY. Do not include any introductory text, explanations, or markdown formatting outside of the JSON structure itself.

    The JSON structure must be as follows:
    {
      "executive_summary": "A concise, 2-3 sentence summary of the overall business performance based on the document.",
      "key_metrics": [
        {
          "metric": "Metric Name (e.g., Total Revenue, Net Profit, YoY Growth)",
          "value": "The value of the metric (e.g., '$150 Million', '12.5%')",
          "commentary": "A brief, one-sentence comment on the significance of this metric."
        }
      ]
    }

    Extract at least 5-7 of the most important metrics you can find in the text.
    """

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Here is the document text:\n\n{text_content}"}
            ]
        )
        json_response = response.choices[0].message.content
        print("Received structured JSON data from Groq LLM.")
        return json.loads(json_response)
    except Exception as e:
        print(f"An error occurred while communicating with the Groq API: {e}")
        return None

def get_insights_chunked(text_content):
    """Process large documents using chunking strategy"""
    
    print("Chunking document for processing...")
    chunks = chunk_text(text_content, max_tokens=6000)
    print(f"Created {len(chunks)} chunks")
    
    # Analyze each chunk
    chunk_results = []
    for i, chunk in enumerate(chunks):
        print(f"Analyzing chunk {i+1}/{len(chunks)}...")
        result = analyze_chunk(chunk, i)
        chunk_results.append(result)
    
    print("Synthesizing final report...")
    final_report = synthesize_final_report(chunk_results, len(chunks))
    
    print("Large document analysis completed!")
    return final_report

def create_excel_report(data, output_path):
    """Creates an Excel report from the extracted data."""
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    df = pd.DataFrame(data['key_metrics'])
    df = df[['metric', 'value', 'commentary']]
    df.to_excel(output_path, index=False)
    print(f"Excel report saved to {output_path}")

def create_word_report(data, output_path):
    """Creates a Word document report from the extracted data."""
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc = Document()
    doc.add_heading('Business Performance Analysis Report', level=1)
    
    doc.add_heading('Executive Summary', level=2)
    summary = data.get('executive_summary', 'No summary available.')
    doc.add_paragraph(summary)
    
    doc.add_heading('Key Metrics', level=2)
    metrics = data.get('key_metrics', [])
    
    if metrics:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Commentary'
        
        for item in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('metric', '')
            row_cells[1].text = item.get('value', '')
            row_cells[2].text = item.get('commentary', '')
    else:
        doc.add_paragraph("No key metrics were extracted.")
        
    doc.save(output_path)
    print(f"Word report saved to {output_path}")

def process_document(file_path):
    """Main orchestration function that handles both PDF and Excel files."""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            # Process PDF
            document_text = extract_text_from_pdf(file_path)
            if not document_text:
                raise ValueError("Could not extract text from the PDF.")

            insights_data = get_insights_from_llm(document_text)
            
        elif file_extension in ['.xlsx', '.xls']:
            # Process Excel
            excel_summary = analyze_excel_file(file_path)
            if not excel_summary:
                raise ValueError("Could not analyze the Excel file.")

            insights_data = get_excel_insights_from_llm(excel_summary)
            
        else:
            raise ValueError(f"Unsupported file type: {file_extension}. Please use PDF, XLSX, or XLS files.")
        
        if not insights_data:
            raise ValueError("Failed to get insights from the language model.")
        
        base_filename = os.path.splitext(os.path.basename(file_path))[0]
        excel_path = os.path.join("outputs", f"{base_filename}_report.xlsx")
        word_path = os.path.join("outputs", f"{base_filename}_report.docx")

        create_excel_report(insights_data, excel_path)
        create_word_report(insights_data, word_path)

        return {
            "success": True,
            "data": insights_data,
            "excel_path": excel_path,
            "word_path": word_path
        }
    except Exception as e:
        print(f"An error occurred during document processing: {e}")
        return {"success": False, "error": str(e)}